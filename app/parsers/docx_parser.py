"""DOCX parser implementation."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from app.models.schemas import Block, BlockType, ParsedDocument
from app.parsers.base import BaseParser


class DocxParser(BaseParser):
    """Parse .docx into paragraph and table-cell blocks."""

    def parse(self, file_path: Path, doc_id: str | None = None) -> ParsedDocument:
        document = Document(file_path)
        blocks: list[Block] = []
        idx = 0

        for p_idx, para in enumerate(document.paragraphs):
            text = (para.text or "").strip()
            if not text:
                continue
            block_type = BlockType.HEADING if para.style and "Heading" in para.style.name else BlockType.PARAGRAPH
            blocks.append(
                Block(
                    block_id=f"p-{idx}",
                    block_type=block_type,
                    text=text,
                    metadata={"paragraph_index": p_idx},
                )
            )
            idx += 1

        for t_idx, table in enumerate(document.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = (cell.text or "").strip()
                    if not text:
                        continue
                    blocks.append(
                        Block(
                            block_id=f"t-{idx}",
                            block_type=BlockType.TABLE_CELL,
                            text=text,
                            metadata={"table_index": t_idx, "row": r_idx, "col": c_idx},
                        )
                    )
                    idx += 1

        return ParsedDocument(
            doc_id=doc_id or file_path.stem,
            source_path=file_path,
            doc_type="docx",
            blocks=blocks,
        )
